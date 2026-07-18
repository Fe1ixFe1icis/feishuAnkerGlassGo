"""Build the submission DOCX by extending the V2 proposal."""

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "Anker_GlassGo产品提案V2_深度优化版.docx"
DST = ROOT / "submission" / "Anker_GlassGo_产品提案_提交版.docx"
ASSETS = ROOT / "submission" / "assets"


def make_para(doc, text, style="Normal"):
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    return p._element


def add_cover_and_elevator(doc):
    """Insert cover + elevator pitch at the very beginning."""
    first_p = doc.paragraphs[0]
    elements = []

    # Spacer
    elements.append(make_para(doc, ""))

    # Read elevator pitch (skip the H1)
    pitch_lines = (ASSETS / "elevator_pitch.md").read_text(encoding="utf-8").splitlines()
    content_lines = pitch_lines[2:] if pitch_lines and pitch_lines[0].startswith("#") else pitch_lines
    for line in reversed(content_lines):
        if line.strip():
            if line.startswith("- **"):
                p = doc.add_paragraph()
                p.add_run("• " + line.replace("- ", "", 1))
                elements.append(p._element)
            else:
                elements.append(make_para(doc, line))
        else:
            elements.append(make_para(doc, ""))

    elements.append(make_para(doc, ""))
    heading = doc.add_paragraph()
    run = heading.add_run("电梯演讲")
    run.bold = True
    run.font.size = Pt(18)
    elements.append(heading._element)
    elements.append(make_para(doc, ""))

    # Cover subtitle
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("AI 原生产品定义工作流 · 提交版").italic = True
    elements.append(sub._element)

    # Cover title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Anker GlassGo")
    run.bold = True
    run.font.size = Pt(28)
    elements.append(title._element)

    for el in elements:
        first_p._element.addprevious(el)


def _add_manual_heading(doc, text, size=18):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    return p


def add_methodology_comparison(doc):
    _add_manual_heading(doc, "方法论对比：AI 原生 vs 经验驱动")
    lines = (ASSETS / "methodology_comparison.md").read_text(encoding="utf-8").splitlines()
    for line in lines:
        stripped = line.strip()
        if not stripped or stripped.startswith("#") or stripped.startswith("|---"):
            continue
        if stripped.startswith("|"):
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            if len(cells) >= 3 and cells[0] != "维度":
                p = doc.add_paragraph()
                run = p.add_run(f"• {cells[0]}：")
                run.bold = True
                p.add_run(f"传统方式——{cells[1]}；AI 原生方式——{cells[2]}。")
        else:
            doc.add_paragraph(stripped)


def add_evidence_sources(doc):
    _add_manual_heading(doc, "证据来源与假设声明")
    lines = (ASSETS / "evidence_sources.md").read_text(encoding="utf-8").splitlines()
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("#"):
            _add_manual_heading(doc, stripped.lstrip("# "), size=14)
            continue
        if stripped.startswith("|---"):
            continue
        if stripped.startswith("|"):
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            if len(cells) >= 5 and cells[0] != "编号":
                p = doc.add_paragraph()
                run = p.add_run(f"{cells[0]}. {cells[2]} — {cells[1]}")
                run.bold = True
                doc.add_paragraph(
                    f"来源类型：{cells[3]}，日期：{cells[4]}，链接：{cells[5] if len(cells) > 5 else '⏳'}"
                )
        else:
            doc.add_paragraph(stripped)


def add_code_appendix(doc):
    _add_manual_heading(doc, "附录：代码仓库与运行方式")
    p = doc.add_paragraph()
    run = p.add_run("本项目已提交完整可运行代码，核心工作流使用 LangGraph 实现六阶段 × 七角色对抗式多 Agent 系统。")
    run.bold = True
    doc.add_paragraph("代码入口：")
    doc.add_paragraph("• py scripts/run_cli.py --auto")
    doc.add_paragraph("• py -m glassgo_workflow.main --auto")
    doc.add_paragraph("• py -m streamlit run streamlit_app.py")
    doc.add_paragraph("• pytest tests/")
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("交互式 Demo 页面：")
    run.bold = True
    doc.add_paragraph("• index.html — 产品提案静态展示")
    doc.add_paragraph("• workflow-demo.html — 六阶段工作流交互演示")


def main():
    DST.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SRC, DST)
    doc = Document(DST)

    add_cover_and_elevator(doc)
    doc.add_page_break()
    add_methodology_comparison(doc)
    doc.add_page_break()
    add_evidence_sources(doc)
    doc.add_page_break()
    add_code_appendix(doc)

    doc.save(DST)
    print(f"Submission DOCX saved to: {DST}")


if __name__ == "__main__":
    main()
